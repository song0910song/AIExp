"""Safe, local document extraction for the project knowledge base."""

from __future__ import annotations

import hashlib
import time
from dataclasses import dataclass
from pathlib import Path

import requests

from .config import PROJECT_ROOT, Settings


class DocumentLoadError(ValueError):
    pass


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    source_path: Path
    source_name: str
    content: str
    sha256: str
    page_count: int | None = None


class PaddleOCRClient:
    """Minimal polling client for a configured PaddleOCR job endpoint.

    Deployments can supply a compatible endpoint through `PADDLEOCR_API_URL`.
    Text PDFs never make a remote call; this client is only a scan fallback.
    """

    def __init__(self, settings: Settings | None = None, session: requests.Session | None = None) -> None:
        self.settings = settings or Settings()
        self.session = session or requests.Session()

    def extract_pdf(self, path: Path) -> str:
        try:
            with path.open("rb") as file_handle:
                response = self.session.post(
                    self.settings.paddleocr_api_url,
                    files={"file": (path.name, file_handle, "application/pdf")},
                    data={"model": self.settings.paddleocr_model},
                    timeout=self.settings.paddleocr_timeout_seconds,
                )
            response.raise_for_status()
            payload = response.json()
        except (OSError, requests.RequestException, ValueError) as error:
            raise DocumentLoadError(f"PaddleOCR submission failed: {error}") from error
        job_id = _nested_value(payload, "job_id") or _nested_value(payload, "id")
        text = _extract_text(payload)
        if text:
            return text
        if not job_id:
            raise DocumentLoadError("PaddleOCR response contained neither text nor a job identifier")
        return self._poll(str(job_id))

    def _poll(self, job_id: str) -> str:
        deadline = time.monotonic() + self.settings.paddleocr_timeout_seconds
        job_url = f"{self.settings.paddleocr_api_url.rstrip('/')}/{job_id}"
        while time.monotonic() < deadline:
            try:
                response = self.session.get(job_url, timeout=min(30, self.settings.paddleocr_timeout_seconds))
                response.raise_for_status()
                payload = response.json()
            except (requests.RequestException, ValueError) as error:
                raise DocumentLoadError(f"PaddleOCR status request failed: {error}") from error
            text = _extract_text(payload)
            if text:
                return text
            status = str(_nested_value(payload, "status") or "").casefold()
            if status in {"failed", "error", "cancelled"}:
                raise DocumentLoadError(f"PaddleOCR job {job_id} ended with status {status}")
            time.sleep(self.settings.paddleocr_poll_interval_seconds)
        raise DocumentLoadError(f"PaddleOCR job {job_id} timed out")


def _checked_path(file_path: str | Path, allowed_root: Path = PROJECT_ROOT) -> Path:
    path = Path(file_path).expanduser().resolve()
    try:
        path.relative_to(allowed_root.resolve())
    except ValueError as error:
        raise DocumentLoadError(f"Document must be within {allowed_root}") from error
    if not path.is_file():
        raise DocumentLoadError(f"Document does not exist or is not a file: {path}")
    return path


def load_document(
    file_path: str | Path,
    *,
    allowed_root: Path = PROJECT_ROOT,
    ocr_client: PaddleOCRClient | None = None,
) -> ParsedDocument:
    """Read an approved project document without executing embedded content."""

    path = _checked_path(file_path, allowed_root)
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        content, page_count = path.read_text(encoding="utf-8"), None
    elif suffix == ".docx":
        content, page_count = _read_docx(path)
    elif suffix == ".pdf":
        content, page_count = _read_pdf(path)
        if not content.strip():
            content = (ocr_client or PaddleOCRClient()).extract_pdf(path)
    else:
        raise DocumentLoadError("Supported document types are .pdf, .docx, .md and .txt")
    content = content.strip()
    if not content:
        raise DocumentLoadError(
            "No extractable text was found. For a scanned PDF, run it through the configured OCR service first."
        )
    return ParsedDocument(
        source_path=path,
        source_name=path.name,
        content=content,
        sha256=hashlib.sha256(path.read_bytes()).hexdigest(),
        page_count=page_count,
    )


def _read_docx(path: Path) -> tuple[str, None]:
    try:
        from docx import Document
    except ImportError as error:  # pragma: no cover - installation issue
        raise DocumentLoadError("python-docx is required to read .docx files") from error
    document = Document(path)
    sections = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                sections.append(" | ".join(cells))
    return "\n\n".join(sections), None


def _read_pdf(path: Path) -> tuple[str, int]:
    try:
        import fitz  # PyMuPDF
    except ImportError as error:  # pragma: no cover - installation issue
        raise DocumentLoadError("PyMuPDF is required to read .pdf files") from error
    pdf = fitz.open(path)
    try:
        pages = [page.get_text("text").strip() for page in pdf]
        return "\n\n".join(text for text in pages if text), len(pdf)
    finally:
        pdf.close()


def _nested_value(payload: object, key: str) -> object | None:
    if isinstance(payload, dict):
        if key in payload:
            return payload[key]
        for value in payload.values():
            nested = _nested_value(value, key)
            if nested is not None:
                return nested
    if isinstance(payload, list):
        for value in payload:
            nested = _nested_value(value, key)
            if nested is not None:
                return nested
    return None


def _extract_text(payload: object) -> str | None:
    """Accept common OCR payload keys without claiming a proprietary schema."""

    for key in ("markdown", "text", "content"):
        value = _nested_value(payload, key)
        if isinstance(value, str) and value.strip():
            return value.strip()
    return None
